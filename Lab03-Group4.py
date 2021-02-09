#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Feb  8 16:49:04 2021

@author: umreenimam
"""

import os
import math
import numpy as np 
import pandas as pd
from scipy.io import loadmat
import matplotlib.pyplot as plt
import seaborn as sns
from pyeeg import bin_power
from docx import Document
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix, accuracy_score
from sklearn.neighbors import KNeighborsClassifier

"""
Step 1: Load the Matlab data files and transpose it
"""
# Setting file names
os.chdir('/Users/umreenimam/documents/masters/masters_classes/cs_5310/week_3/lab_chapter3')
filename = 'Acquisition-15-data.mat'
prefile = 'Pre.mat'
medfile = 'Med.mat'
postfile = 'Post.mat'

# Loading data 
data = loadmat(filename)
predata = loadmat(prefile)
meddata = loadmat(medfile)
postdata = loadmat(postfile)
eeg = data['data']
pre = predata['data']
med = meddata['data']
post = postdata['data']

# Transpose data and create dataframes 
eeg_transposed = pd.DataFrame(eeg.T)
pre_transposed = pd.DataFrame(pre.T)
med_transposed = pd.DataFrame(med.T)
post_transposed = pd.DataFrame(post.T)

"""
Step 2 through 4: Run EEG calculations to build a data frame with alpha values
"""
# Computing Alpha PSI for 3 dataframes 
# Function to compute alpha PSI
def getPSI(data):
    band = [0.5, 4, 7, 12, 30]
    fs = 1024
    sample = 1024
    cols = data.shape[1]
    rows = math.floor(data.shape[0] / sample)
    alpha_psi = pd.DataFrame(0, index=range(rows), columns=range(cols))
    for x in range(cols):
        bin_powers = []
        for y in range(rows):
            psis = bin_power(data.iloc[(y * sample):((y + 1) * sample), x], band, fs)[0][2]
            bin_powers.append(psis)
        alpha_psi[x] = bin_powers
    return alpha_psi

# Running function for each dataset
prePSI = getPSI(pre_transposed)
medPSI = getPSI(med_transposed)
postPSI = getPSI(post_transposed)

# Brain state labels 
pre_label = ["Pre"] * prePSI.shape[0]
med_label = ["Med"] * medPSI.shape[0]
post_label = ["Post"] * postPSI.shape[0]

# Combine dataframes
frames = [prePSI, medPSI, postPSI]
combinedPSI = pd.concat(frames, ignore_index=True)

# Combine label lists
combinedLabels = pre_label + med_label + post_label

"""
Step 5 through 7: Remove any co-linear columns
"""
# Create correlation coefficient matrix 
corrMat = combinedPSI.corr()
rows = corrMat.shape[0]
cols = corrMat.shape[1]

# Initialize an empty list to keep track of which columns need to be removed
columns_to_remove = []

for i in range(rows):
    for j in range(cols):
        if i == j:
            # If the row number equals the column number, then exit this for
            # loop because we don't want to accidentally remove the same
            # column number twice.
            break
        if corrMat.iloc[i, j] > 0.9 or corrMat.iloc[i, j] < -0.9:
            # If any values in the correlation matrix are greater than 0.9
            # or less than -0.9, add their index number to the list of columns
            # that need to be removed.
            columns_to_remove.append(i)
            break

# Remove any columns that had a correlation factor that was too high
alpha_df = combinedPSI.drop(columns = columns_to_remove, axis = 1)

# Create a new correlation matrix and print it
alpha_corr_2 = alpha_df.corr()
print(alpha_corr_2)


sns.set_theme(style = "white")

plt.figure(figsize=(6,5))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(corrMat, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.1,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('fig1.png')
plt.show


plt.figure(figsize=(6,5))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(alpha_corr_2, annot = True, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.5,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('fig2.png')
plt.show

"""
Step 8 and 9: Normalize the data and split data into train-test data sets
"""
# Normalize data in every remaining channel using the min-max normalization
scaler = MinMaxScaler()
normalized_corr = scaler.fit_transform(alpha_df)
normalized_corr = pd.DataFrame(normalized_corr)

print(normalized_corr)

X = normalized_corr
y = combinedLabels

X_train, X_test, labels_train, labels_test = train_test_split(
   X, 
   y, 
   random_state = 42, 
   test_size = 0.2, 
   stratify = combinedLabels)

"""
Step 10 through 12: Build KNN model and generate confusion matrix
"""
# Train the kNN model using the training dataset
n_neighbors = round(math.sqrt(X_train.shape[1]))
classifier = KNeighborsClassifier(n_neighbors = n_neighbors)
classifier.fit(X_train, labels_train)

# Use the kNN model to predict the target feature of the testing dataset
prediction = classifier.predict(X_test)
pd.crosstab(labels_test, prediction)

# Create a confusion matrix to compare the predicted state activities to
# the actual activities and compute the accuracy
confusion_mat = confusion_matrix(labels_test, prediction)
print(confusion_mat)

# Compute the accuracy
accuracy_rate = accuracy_score(labels_test, prediction)
print(accuracy_rate)

"""
Generate a Word document with the results from this lab
"""

doc = Document()

doc.add_heading('Correlation Coefficient Matrix Before Removing Co-linearity', level = 1)

doc.add_picture('fig1.png')
doc.add_paragraph()
doc.add_page_break()

doc.add_heading('Correlation Coefficient Matrix After Removing Co-linearity', level = 1)

doc.add_picture('fig2.png')
doc.add_paragraph()

doc.add_heading('Confusion Matrix:', level = 1)

table = doc.add_table(rows = confusion_mat.shape[0] + 1, cols = confusion_mat.shape[1] + 1)
table.style = 'Medium Grid 3 Accent 3'

row = table.rows[0]
row.cells[1].text = 'Predicted Med'
row.cells[2].text = 'Predicted Post'
row.cells[3].text = 'Predicted Pre'

col = table.columns[0]
col.cells[1].text = 'Actual Med'
col.cells[2].text = 'Actual Post'
col.cells[3].text = 'Actual Pre'

for i in range(confusion_mat.shape[0]):
    for j in range(confusion_mat.shape[1]):
        table.cell(i + 1, j + 1).text = str(confusion_mat[i, j])

doc.add_paragraph()
doc.add_paragraph('Accuracy Rate: {}%'.format(round(accuracy_rate * 100, 1)))
doc.save('Chapter3-Lab-Group4.docx')